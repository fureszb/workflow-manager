from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, File
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from sqlalchemy import or_
from typing import List, Optional
import os
import uuid
import shutil
import io
import json

from app.core.database import get_db
from app.models.models import Document, DocumentChunk
from app.schemas.schemas import DocumentResponse, DocumentUpdate, DocumentPreviewResponse, DocumentSearchResult, DocumentSummaryResponse
from app.routers.websocket_router import broadcast_notification

router = APIRouter(prefix="/documents")

# Configure upload directory - /storage/documents/ as per acceptance criteria
STORAGE_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))), "storage", "documents")
os.makedirs(STORAGE_DIR, exist_ok=True)

# Versions directory for storing old file versions
VERSIONS_DIR = os.path.join(STORAGE_DIR, "versions")
os.makedirs(VERSIONS_DIR, exist_ok=True)

# Maximum number of versions to keep (excluding current)
MAX_VERSIONS = 2

# Allowed file extensions
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".txt"}


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text content from a PDF file."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return "\n\n".join(text_parts)
    except Exception as e:
        return f"[Hiba a PDF olvasásakor: {str(e)}]"


def extract_text_from_docx(file_path: str) -> str:
    """Extract text content from a DOCX file."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        return "\n\n".join(paragraphs)
    except Exception as e:
        return f"[Hiba a DOCX olvasásakor: {str(e)}]"


def extract_text_from_xlsx(file_path: str) -> tuple:
    """Extract content from an XLSX file as both text and table data."""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(file_path, data_only=True)

        text_parts = []
        all_table_data = []

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_parts.append(f"--- Munkalap: {sheet_name} ---")

            sheet_data = []
            for row in sheet.iter_rows(values_only=True):
                row_values = [str(cell) if cell is not None else "" for cell in row]
                if any(v.strip() for v in row_values):
                    sheet_data.append(row_values)
                    text_parts.append("\t".join(row_values))

            if sheet_data:
                all_table_data.extend(sheet_data)

        return "\n".join(text_parts), all_table_data
    except Exception as e:
        return f"[Hiba az XLSX olvasásakor: {str(e)}]", None


def extract_text_from_txt(file_path: str) -> str:
    """Read text content from a TXT file."""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        try:
            with open(file_path, "r", encoding="latin-1") as f:
                return f.read()
        except Exception as e:
            return f"[Hiba a TXT olvasásakor: {str(e)}]"
    except Exception as e:
        return f"[Hiba a TXT olvasásakor: {str(e)}]"


def extract_document_content(file_path: str, file_type: str) -> tuple:
    """Extract content from a document based on its type.

    Returns (text_content, table_data) where table_data is only set for XLSX files.
    """
    if not os.path.exists(file_path):
        return None, None

    file_type = file_type.lower() if file_type else ""

    if file_type == "pdf":
        return extract_text_from_pdf(file_path), None
    elif file_type == "docx":
        return extract_text_from_docx(file_path), None
    elif file_type == "xlsx":
        return extract_text_from_xlsx(file_path)
    elif file_type == "txt":
        return extract_text_from_txt(file_path), None
    else:
        return None, None


def highlight_search_matches(text: str, query: str) -> list:
    """Find and highlight search matches in text.

    Returns a list of SearchMatch objects with line numbers and highlighted text.
    """
    if not text or not query:
        return []

    query_lower = query.lower()
    lines = text.split("\n")
    matches = []

    for line_num, line in enumerate(lines, start=1):
        if query_lower in line.lower():
            # Find all match positions and create highlighted version
            highlighted = line
            start_idx = 0
            while True:
                pos = highlighted.lower().find(query_lower, start_idx)
                if pos == -1:
                    break
                # Insert highlight markers
                match_text = highlighted[pos:pos + len(query)]
                highlighted = highlighted[:pos] + f"**{match_text}**" + highlighted[pos + len(query):]
                start_idx = pos + len(query) + 4  # Account for added markers

            matches.append({
                "line_number": line_num,
                "text": line.strip(),
                "highlighted_text": highlighted.strip()
            })

    return matches


def get_file_extension(filename: str) -> str:
    """Extract file extension from filename."""
    if not filename:
        return ""
    return os.path.splitext(filename)[1].lower()


def is_allowed_file(filename: str) -> bool:
    """Check if file extension is allowed."""
    ext = get_file_extension(filename)
    return ext in ALLOWED_EXTENSIONS


def _archive_old_version(existing_doc: Document, db: Session):
    """Move an existing document to the versions directory.

    Enforces max 2 versions - deletes oldest if limit exceeded.
    """
    # Find all existing versions of this document chain
    if existing_doc.parent_id:
        # Document is already a version, find parent
        parent_id = existing_doc.parent_id
    else:
        parent_id = existing_doc.id

    # Get all versions for this document chain (excluding the current one being replaced)
    versions = db.query(Document).filter(
        (Document.parent_id == parent_id) | (Document.id == parent_id),
        Document.id != existing_doc.id
    ).order_by(Document.version.desc()).all()

    # If we already have MAX_VERSIONS, delete the oldest
    if len(versions) >= MAX_VERSIONS:
        # Sort by version to find oldest
        versions_sorted = sorted(versions, key=lambda d: d.version)
        oldest = versions_sorted[0]

        # Delete oldest version file from disk
        if oldest.file_path and os.path.exists(oldest.file_path):
            os.remove(oldest.file_path)

        # Delete from database
        db.delete(oldest)
        db.commit()

    # Move the current document to versions directory
    if existing_doc.file_path and os.path.exists(existing_doc.file_path):
        # Create versioned filename
        version_filename = f"{existing_doc.filename}_v{existing_doc.version}"
        version_path = os.path.join(VERSIONS_DIR, version_filename)

        # Move file to versions directory
        shutil.move(existing_doc.file_path, version_path)

        # Update file_path in database
        existing_doc.file_path = version_path
        db.commit()


@router.post("/upload", response_model=DocumentResponse, status_code=201)
async def upload_document(
    file: UploadFile = File(...),
    category: Optional[str] = Query(None),
    is_knowledge_base: bool = Query(False, alias="is_knowledge_base"),
    db: Session = Depends(get_db),
):
    """Upload a document (PDF, DOCX, XLSX, TXT).

    Files are stored in /storage/documents/ with UUID filenames.
    When uploading a file with the same name, the old version is moved
    to /storage/documents/versions/ with max 2 versions retained.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="Nincs fájlnév megadva")

    if not is_allowed_file(file.filename):
        allowed = ", ".join(ALLOWED_EXTENSIONS)
        raise HTTPException(
            status_code=400,
            detail=f"Nem megengedett fájltípus. Engedélyezett típusok: {allowed}"
        )

    # Check if a document with the same original_filename already exists
    existing_doc = db.query(Document).filter(
        Document.original_filename == file.filename,
        Document.parent_id.is_(None)  # Only check current versions (not archived ones)
    ).first()

    parent_id = None
    new_version = 1

    if existing_doc:
        # Archive the old version
        _archive_old_version(existing_doc, db)

        # New document will be a new version
        parent_id = existing_doc.parent_id if existing_doc.parent_id else existing_doc.id
        new_version = existing_doc.version + 1

        # Mark old document as archived (set parent_id to itself if not already set)
        if not existing_doc.parent_id:
            existing_doc.parent_id = existing_doc.id
            db.commit()

    # Generate unique filename
    file_ext = get_file_extension(file.filename)
    unique_filename = f"{uuid.uuid4()}{file_ext}"
    file_path = os.path.join(STORAGE_DIR, unique_filename)

    # Save file to disk (no size limit as per requirements)
    content = await file.read()
    with open(file_path, "wb") as f:
        f.write(content)

    # Create document record
    document = Document(
        filename=unique_filename,
        original_filename=file.filename,
        file_path=file_path,
        file_type=file_ext.lstrip(".") if file_ext else None,
        file_size=len(content),
        category=category if category else (existing_doc.category if existing_doc else None),
        is_knowledge=is_knowledge_base if not existing_doc else existing_doc.is_knowledge,
        version=new_version,
        parent_id=None,  # Current version has no parent
    )
    db.add(document)
    db.commit()
    db.refresh(document)

    return document


@router.get("", response_model=List[DocumentResponse])
def list_documents(
    file_type: Optional[str] = Query(None, description="Filter by file type (pdf, docx, xlsx, txt)"),
    category: Optional[str] = Query(None, description="Filter by category"),
    is_knowledge_base: Optional[bool] = Query(None, alias="is_knowledge_base", description="Filter by knowledge base flag"),
    db: Session = Depends(get_db),
):
    """List documents with optional filtering by type, category, and knowledge base flag."""
    query = db.query(Document)

    if file_type:
        query = query.filter(Document.file_type == file_type.lower())

    if category:
        query = query.filter(Document.category == category)

    if is_knowledge_base is not None:
        query = query.filter(Document.is_knowledge == is_knowledge_base)

    return query.order_by(Document.created_at.desc()).all()


# IMPORTANT: /search must be defined BEFORE /{doc_id} to avoid route conflict
@router.get("/search", response_model=List[DocumentSearchResult])
def search_documents(
    q: str = Query("", description="Search query"),
    content_search: bool = Query(False, description="Search within document content"),
    db: Session = Depends(get_db),
):
    """Search documents by original filename and optionally within content.

    When content_search=true, searches within document text content and returns
    matching lines with highlights.
    """
    from app.schemas.schemas import SearchMatch

    if not q:
        return []

    search_term = f"%{q}%"

    # First, get documents matching filename
    documents = db.query(Document).filter(
        Document.original_filename.ilike(search_term),
        Document.parent_id.is_(None)  # Only current versions
    ).order_by(Document.created_at.desc()).limit(50).all()

    results = []

    if content_search:
        # Search within document content
        all_docs = db.query(Document).filter(
            Document.parent_id.is_(None)  # Only current versions
        ).order_by(Document.created_at.desc()).limit(100).all()

        for doc in all_docs:
            text_content, _ = extract_document_content(doc.file_path, doc.file_type)

            if text_content:
                matches = highlight_search_matches(text_content, q)

                if matches:
                    results.append({
                        "id": doc.id,
                        "original_filename": doc.original_filename,
                        "file_type": doc.file_type,
                        "file_size": doc.file_size,
                        "category": doc.category,
                        "created_at": doc.created_at,
                        "matches": matches[:10],  # Limit matches per document
                        "match_count": len(matches)
                    })

        # Also include filename matches that weren't found in content search
        filename_match_ids = {r["id"] for r in results}
        for doc in documents:
            if doc.id not in filename_match_ids:
                results.append({
                    "id": doc.id,
                    "original_filename": doc.original_filename,
                    "file_type": doc.file_type,
                    "file_size": doc.file_size,
                    "category": doc.category,
                    "created_at": doc.created_at,
                    "matches": [{
                        "line_number": 0,
                        "text": f"Fájlnév találat: {doc.original_filename}",
                        "highlighted_text": doc.original_filename.replace(
                            q, f"**{q}**"
                        ) if q.lower() in doc.original_filename.lower() else doc.original_filename
                    }],
                    "match_count": 1
                })
    else:
        # Filename-only search (original behavior)
        for doc in documents:
            results.append({
                "id": doc.id,
                "original_filename": doc.original_filename,
                "file_type": doc.file_type,
                "file_size": doc.file_size,
                "category": doc.category,
                "created_at": doc.created_at,
                "matches": [],
                "match_count": 0
            })

    return results[:50]  # Limit total results


@router.get("/categories", response_model=List[str])
def list_categories(db: Session = Depends(get_db)):
    """List all unique document categories.

    Returns a sorted list of all non-null categories used in documents.
    Categories are created dynamically when uploading or updating documents.
    """
    categories = db.query(Document.category).filter(
        Document.category.isnot(None),
        Document.category != ""
    ).distinct().all()

    # Extract strings from tuples and sort
    return sorted([cat[0] for cat in categories])


@router.get("/{doc_id}/preview", response_model=DocumentPreviewResponse)
def preview_document(doc_id: int, db: Session = Depends(get_db)):
    """Get document preview content.

    - PDF: Returns preview_url for inline display via iframe/embed
    - DOCX/TXT: Returns extracted text content
    - XLSX: Returns both text content and structured table_data
    """
    document = db.query(Document).filter(Document.id == doc_id).first()

    if not document:
        raise HTTPException(status_code=404, detail="Dokumentum nem található")

    if not os.path.exists(document.file_path):
        return DocumentPreviewResponse(
            id=document.id,
            original_filename=document.original_filename,
            file_type=document.file_type,
            error="Fájl nem található a lemezen"
        )

    file_type = document.file_type.lower() if document.file_type else ""

    # For PDF, provide URL for inline viewing
    if file_type == "pdf":
        return DocumentPreviewResponse(
            id=document.id,
            original_filename=document.original_filename,
            file_type=document.file_type,
            preview_url=f"/api/v1/documents/{doc_id}/raw"
        )

    # For other types, extract content
    text_content, table_data = extract_document_content(document.file_path, file_type)

    if text_content and text_content.startswith("[Hiba"):
        return DocumentPreviewResponse(
            id=document.id,
            original_filename=document.original_filename,
            file_type=document.file_type,
            error=text_content
        )

    return DocumentPreviewResponse(
        id=document.id,
        original_filename=document.original_filename,
        file_type=document.file_type,
        content=text_content,
        table_data=table_data
    )


@router.get("/{doc_id}/raw")
def get_raw_document(doc_id: int, db: Session = Depends(get_db)):
    """Serve document file for inline viewing (especially PDF).

    Returns the file with appropriate Content-Type for browser rendering.
    """
    document = db.query(Document).filter(Document.id == doc_id).first()

    if not document:
        raise HTTPException(status_code=404, detail="Dokumentum nem található")

    if not os.path.exists(document.file_path):
        raise HTTPException(status_code=404, detail="Fájl nem található a lemezen")

    # Map file types to MIME types
    mime_types = {
        "pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "txt": "text/plain; charset=utf-8",
    }

    file_type = document.file_type.lower() if document.file_type else "application/octet-stream"
    media_type = mime_types.get(file_type, "application/octet-stream")

    return FileResponse(
        path=document.file_path,
        media_type=media_type,
        filename=document.original_filename,
        headers={
            "Content-Disposition": f'inline; filename="{document.original_filename}"'
        }
    )


@router.get("/{doc_id}", response_model=DocumentResponse)
def get_document(doc_id: int, db: Session = Depends(get_db)):
    """Get document details by ID."""
    document = db.query(Document).filter(Document.id == doc_id).first()

    if not document:
        raise HTTPException(status_code=404, detail="Dokumentum nem található")

    return document


@router.get("/{doc_id}/download")
def download_document(doc_id: int, db: Session = Depends(get_db)):
    """Download a document by ID."""
    document = db.query(Document).filter(Document.id == doc_id).first()

    if not document:
        raise HTTPException(status_code=404, detail="Dokumentum nem található")

    if not os.path.exists(document.file_path):
        raise HTTPException(status_code=404, detail="Fájl nem található a lemezen")

    return FileResponse(
        path=document.file_path,
        filename=document.original_filename,
        media_type="application/octet-stream",
    )


@router.get("/{doc_id}/versions", response_model=List[DocumentResponse])
def get_versions(doc_id: int, db: Session = Depends(get_db)):
    """Get version history of a document.

    Returns all versions including the current document and archived versions.
    Versions are ordered by version number descending (newest first).
    """
    document = db.query(Document).filter(Document.id == doc_id).first()

    if not document:
        raise HTTPException(status_code=404, detail="Dokumentum nem található")

    # Find all documents with the same original_filename
    versions = db.query(Document).filter(
        Document.original_filename == document.original_filename
    ).order_by(Document.version.desc()).all()

    return versions


@router.get("/{doc_id}/versions/{version_id}/download")
def download_version(doc_id: int, version_id: int, db: Session = Depends(get_db)):
    """Download a specific version of a document."""
    # Verify the main document exists
    document = db.query(Document).filter(Document.id == doc_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Dokumentum nem található")

    # Get the requested version
    version = db.query(Document).filter(Document.id == version_id).first()
    if not version:
        raise HTTPException(status_code=404, detail="Verzió nem található")

    # Verify the version belongs to the same document chain (same original_filename)
    if version.original_filename != document.original_filename:
        raise HTTPException(status_code=400, detail="A verzió nem tartozik ehhez a dokumentumhoz")

    if not os.path.exists(version.file_path):
        raise HTTPException(status_code=404, detail="Fájl nem található a lemezen")

    return FileResponse(
        path=version.file_path,
        filename=f"{version.original_filename}_v{version.version}",
        media_type="application/octet-stream",
    )


@router.delete("/{doc_id}")
def delete_document(doc_id: int, db: Session = Depends(get_db)):
    """Delete a document by ID."""
    document = db.query(Document).filter(Document.id == doc_id).first()

    if not document:
        raise HTTPException(status_code=404, detail="Dokumentum nem található")

    # Delete physical file if it exists
    if os.path.exists(document.file_path):
        os.remove(document.file_path)

    # Delete database record
    db.delete(document)
    db.commit()

    return {"message": "Dokumentum törölve"}


@router.put("/{doc_id}", response_model=DocumentResponse)
def update_document(
    doc_id: int,
    payload: DocumentUpdate,
    db: Session = Depends(get_db),
):
    """Update document metadata (category, is_knowledge)."""
    document = db.query(Document).filter(Document.id == doc_id).first()

    if not document:
        raise HTTPException(status_code=404, detail="Dokumentum nem található")

    for key, value in payload.model_dump(exclude_unset=True).items():
        setattr(document, key, value)

    db.commit()
    db.refresh(document)

    return document


# Note: Chunking and embedding logic has been moved to app.services.rag_service
# for the full RAG pipeline with FAISS integration.


def delete_document_chunks(document_id: int, db: Session):
    """Delete all chunks for a document (when removing from knowledge base)."""
    db.query(DocumentChunk).filter(DocumentChunk.document_id == document_id).delete()
    db.commit()


@router.post("/{doc_id}/toggle-knowledge", response_model=DocumentResponse)
async def toggle_knowledge(doc_id: int, db: Session = Depends(get_db)):
    """Toggle the is_knowledge flag for a document.

    When adding to knowledge base:
    - Extracts text from document
    - Creates chunks using hybrid chunking (semantic + 500 token limit)
    - Generates embeddings with Ollama
    - Updates FAISS index and stores chunks in DocumentChunk table

    When removing from knowledge base:
    - Deletes all document chunks from DB and FAISS index
    """
    from app.services.rag_service import index_document, remove_document_from_index
    from app.models.models import AIKnowledgeLog

    document = db.query(Document).filter(Document.id == doc_id).first()

    if not document:
        raise HTTPException(status_code=404, detail="Dokumentum nem található")

    # Toggle the flag
    new_state = not document.is_knowledge
    document.is_knowledge = new_state

    if new_state:
        # Adding to knowledge base - index document with RAG pipeline
        chunks_created, error = await index_document(document, db)
        if error and chunks_created == 0:
            # Rollback if complete failure
            document.is_knowledge = False
            db.commit()
            # Log failed attempt
            log_entry = AIKnowledgeLog(
                document_id=doc_id,
                action="added",
                chunks_processed=0,
                status="failed",
                error_message=error,
            )
            db.add(log_entry)
            db.commit()
            raise HTTPException(status_code=500, detail=error)

        # Log successful addition
        log_entry = AIKnowledgeLog(
            document_id=doc_id,
            action="added",
            chunks_processed=chunks_created,
            status="completed" if not error else "partial",
            error_message=error,
        )
        db.add(log_entry)
    else:
        # Get chunk count before removal for logging
        chunk_count = db.query(DocumentChunk).filter(DocumentChunk.document_id == doc_id).count()

        # Removing from knowledge base - delete from index
        await remove_document_from_index(doc_id, db)

        # Log removal
        log_entry = AIKnowledgeLog(
            document_id=doc_id,
            action="removed",
            chunks_processed=chunk_count,
            status="completed",
        )
        db.add(log_entry)

    db.commit()
    db.refresh(document)

    # Send notification for document knowledge base processing
    if new_state:
        await broadcast_notification(
            message=f"'{document.original_filename}' hozzáadva a tudásbázishoz",
            level="success",
            title="Dokumentum feldolgozás kész",
            action_url="/knowledge"
        )
    else:
        await broadcast_notification(
            message=f"'{document.original_filename}' eltávolítva a tudásbázisból",
            level="info",
            title="Dokumentum eltávolítva",
            action_url="/knowledge"
        )

    return document


@router.post("/{doc_id}/summarize", response_model=DocumentSummaryResponse)
async def summarize_document(doc_id: int, db: Session = Depends(get_db)):
    """Generate AI summary for a document.

    Uses the configured AI provider (Ollama or OpenRouter) to generate
    a summary of the document content. The summary is stored in the
    document's summary field.
    """
    from app.services.ai_service import get_ai_settings, send_chat_message

    document = db.query(Document).filter(Document.id == doc_id).first()

    if not document:
        raise HTTPException(status_code=404, detail="Dokumentum nem található")

    # Extract document content
    text_content, _ = extract_document_content(document.file_path, document.file_type)

    if not text_content:
        raise HTTPException(
            status_code=400,
            detail="Nem sikerült szöveget kinyerni a dokumentumból"
        )

    if text_content.startswith("[Hiba"):
        raise HTTPException(status_code=400, detail=text_content)

    # Truncate content if too long (limit to ~4000 chars to avoid token limits)
    max_content_length = 4000
    if len(text_content) > max_content_length:
        text_content = text_content[:max_content_length] + "\n\n[... tartalom csonkolva ...]"

    # Build prompt for summarization
    system_prompt = """Te egy segítőkész asszisztens vagy, aki dokumentumok összefoglalását készíti magyar nyelven.
Az összefoglaló legyen:
- Tömör és lényegre törő (maximum 200 szó)
- A legfontosabb információkat emelje ki
- Jól strukturált, könnyen olvasható
- Magyar nyelven íródjon"""

    user_message = f"""Kérlek, készíts egy rövid összefoglalót az alábbi dokumentumról:

Dokumentum neve: {document.original_filename}

Tartalom:
{text_content}

Készítsd el az összefoglalót magyar nyelven!"""

    try:
        # Generate summary using AI
        summary_text, _, _ = await send_chat_message(
            messages=[{"role": "user", "content": user_message}],
            db=db,
            system_prompt=system_prompt,
        )

        # Save summary to document
        document.summary = summary_text
        db.commit()
        db.refresh(document)

        # Send notification
        await broadcast_notification(
            message=f"'{document.original_filename}' összefoglalója elkészült",
            level="success",
            title="Összefoglaló generálva",
            action_url="/documents"
        )

        return DocumentSummaryResponse(
            id=document.id,
            original_filename=document.original_filename,
            summary=summary_text,
            message="Összefoglaló sikeresen generálva"
        )

    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Hiba az összefoglaló generálása közben: {str(e)}"
        )


# --- RAG / Knowledge Base Endpoints ---

@router.get("/rag/stats")
def get_rag_stats(db: Session = Depends(get_db)):
    """Get RAG index statistics.

    Returns information about the FAISS index and indexed documents.
    """
    from app.services.rag_service import get_index_stats

    return get_index_stats(db)


@router.post("/rag/reindex")
async def reindex_knowledge_base(db: Session = Depends(get_db)):
    """Reindex all documents marked as knowledge base.

    This will regenerate chunks and embeddings for all is_knowledge=true documents.
    Useful for rebuilding the index after model changes or data corruption.
    """
    from app.services.rag_service import reindex_all_knowledge_documents

    result = await reindex_all_knowledge_documents(db)
    return result


@router.post("/rag/search")
async def search_knowledge_base(
    query: str = Query(..., description="Search query"),
    k: int = Query(5, description="Number of results to return", ge=1, le=20),
    db: Session = Depends(get_db),
):
    """Search the knowledge base using semantic similarity.

    Uses the FAISS index to find document chunks most similar to the query.
    Returns chunks with their content, source document info, and similarity scores.
    """
    from app.services.rag_service import search_similar_chunks

    if not query.strip():
        return []

    results = await search_similar_chunks(query, db, k)
    return results
